#
#  Copyright 2024 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#
import json
import re
import traceback
import uuid
import base64
import time
import io
from copy import deepcopy

import trio
from flask import Response, jsonify, request
from flask_login import current_user, login_required
from itsdangerous.url_safe import URLSafeTimedSerializer as Serializer

from api import settings
from api.db import LLMType, StatusEnum
from api.db.db_models import APIToken
from api.db.services.conversation_service import ConversationService, structure_answer
from api.db.services.dialog_service import DialogService, ask, chat
from api.db.services.knowledgebase_service import KnowledgebaseService
from api.db.services.llm_service import LLMBundle, TenantService
from api.db.services.user_service import UserTenantService, UserService
from api.db.services.write_service import upload_image, write_dialog
from api.utils.api_utils import get_data_error_result, get_json_result, server_error_response, validate_request
from graphrag.general.mind_map_extractor import MindMapExtractor
from rag.app.tag import label_question
from rag.utils.redis_conn import REDIS_CONN

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import olefile
    from docx2txt import process as docx2txt_process
    DOC_AVAILABLE = True
except ImportError:
    DOC_AVAILABLE = False


def extract_file_content(file_content_bytes, filename, content_type):
    """
    根据文件类型提取文件内容
    
    Args:
        file_content_bytes: 文件的二进制内容
        filename: 文件名
        content_type: MIME类型
    
    Returns:
        str: 提取的文本内容
    """
    try:
        # 获取文件扩展名
        file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
        
        # 处理旧版Word文档(.doc)
        if file_ext == 'doc':
            try:
                # 对于.doc文件，由于格式复杂，我们提供一个友好的提示
                return f'检测到旧版Word文档(.doc格式)：{filename}\n\n由于.doc格式的复杂性，建议您：\n1. 将文件另存为.docx格式后重新上传\n2. 或者复制文档内容直接粘贴到聊天框中\n\n这样可以确保内容被正确解析和处理。'
            except Exception as e:
                print(f".doc文件处理错误: {e}")
                return f'无法处理.doc格式文件：{filename}\n建议转换为.docx格式或直接粘贴文本内容。'
        
        # 处理新版Word文档(.docx)
        elif file_ext == 'docx' or 'word' in content_type.lower():
            if DOCX_AVAILABLE:
                try:
                    # 使用python-docx处理Word文档
                    doc_stream = io.BytesIO(file_content_bytes)
                    doc = Document(doc_stream)
                    
                    # 提取所有段落文本
                    paragraphs = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            paragraphs.append(paragraph.text.strip())
                    
                    # 提取表格内容
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                paragraphs.append(' | '.join(row_text))
                    
                    return '\n\n'.join(paragraphs) if paragraphs else '无法提取文档内容'
                    
                except Exception as e:
                    print(f"Word文档处理错误: {e}")
                    return f'Word文档解析失败：{filename}\n建议检查文件是否损坏或转换为文本格式。'
            else:
                return f'缺少Word文档处理库，无法解析：{filename}\n建议将内容复制粘贴到聊天框中。'
        
        # 处理文本文件
        elif file_ext in ['txt', 'md', 'py', 'js', 'html', 'css', 'json', 'xml', 'csv'] or 'text' in content_type.lower():
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'big5', 'latin1']
            for encoding in encodings:
                try:
                    return file_content_bytes.decode(encoding)
                except UnicodeDecodeError:
                    continue
            # 如果所有编码都失败，使用错误忽略模式
            return file_content_bytes.decode('utf-8', errors='ignore')
        
        # 处理PDF文件
        elif file_ext == 'pdf':
            return f'暂不支持PDF文件内容提取：{filename}\n建议转换为Word或文本格式后重新上传。'
        
        # 其他文件类型
        else:
            # 尝试作为文本文件处理
            try:
                # 先尝试UTF-8
                text_content = file_content_bytes.decode('utf-8')
                # 检查是否包含过多的非打印字符（可能是二进制文件）
                non_printable_ratio = sum(1 for c in text_content if ord(c) < 32 and c not in '\n\r\t') / len(text_content) if text_content else 0
                if non_printable_ratio > 0.3:  # 如果超过30%是非打印字符，可能是二进制文件
                    return f'检测到二进制文件：{filename}\n文件类型：{content_type}\n建议上传文本格式的文件以获得更好的处理效果。'
                return text_content
            except UnicodeDecodeError:
                # 尝试其他编码
                encodings = ['gbk', 'gb2312', 'big5', 'latin1']
                for encoding in encodings:
                    try:
                        return file_content_bytes.decode(encoding, errors='ignore')
                    except:
                        continue
                return f'无法解析文件内容：{filename}\n文件类型：{content_type}\n建议转换为支持的格式（.txt, .docx等）。'
                
    except Exception as e:
        print(f"文件内容提取错误: {e}")
        return f'文件处理失败：{filename}\n错误信息：{str(e)}\n建议检查文件格式或重新上传。'


@manager.route("/set", methods=["POST"])  # type: ignore # noqa: F821
@login_required
def set_conversation():
    req = request.json
    conv_id = req.get("conversation_id")
    is_new = req.get("is_new")
    del req["is_new"]
    if not is_new:
        del req["conversation_id"]
        try:
            if not ConversationService.update_by_id(conv_id, req):
                return get_data_error_result(message="Conversation not found!")
            e, conv = ConversationService.get_by_id(conv_id)
            if not e:
                return get_data_error_result(message="Fail to update a conversation!")
            conv = conv.to_dict()
            return get_json_result(data=conv)
        except Exception as e:
            return server_error_response(e)

    try:
        e, dia = DialogService.get_by_id(req["dialog_id"])
        if not e:
            return get_data_error_result(message="Dialog not found")
        conv = {"id": conv_id, "dialog_id": req["dialog_id"], "name": req.get("name", "New conversation"), "message": [{"role": "assistant", "content": dia.prompt_config["prologue"]}]}
        ConversationService.save(**conv)
        return get_json_result(data=conv)
    except Exception as e:
        return server_error_response(e)


@manager.route("/get", methods=["GET"])  # type: ignore # type: ignore # noqa: F821
@login_required
def get():
    conv_id = request.args["conversation_id"]
    try:
        e, conv = ConversationService.get_by_id(conv_id)
        if not e:
            return get_data_error_result(message="Conversation not found!")
        tenants = UserTenantService.query(user_id=current_user.id)
        avatar = None
        for tenant in tenants:
            dialog = DialogService.query(tenant_id=tenant.tenant_id, id=conv.dialog_id)
            if dialog and len(dialog) > 0:
                avatar = dialog[0].icon
                break
        else:
            return get_json_result(data=False, message="Only owner of conversation authorized for this operation.", code=settings.RetCode.OPERATING_ERROR)

        def get_value(d, k1, k2):
            return d.get(k1, d.get(k2))

        for ref in conv.reference:
            if isinstance(ref, list):
                continue
            ref["chunks"] = [
                {
                    "id": get_value(ck, "chunk_id", "id"),
                    "content": get_value(ck, "content", "content_with_weight"),
                    "document_id": get_value(ck, "doc_id", "document_id"),
                    "document_name": get_value(ck, "docnm_kwd", "document_name"),
                    "dataset_id": get_value(ck, "kb_id", "dataset_id"),
                    "image_id": get_value(ck, "image_id", "img_id"),
                    "positions": get_value(ck, "positions", "position_int"),
                }
                for ck in ref.get("chunks", [])
            ]

        conv = conv.to_dict()
        conv["avatar"] = avatar
        return get_json_result(data=conv)
    except Exception as e:
        return server_error_response(e)


@manager.route("/getsse/<dialog_id>", methods=["GET"])  # type: ignore # noqa: F821
def getsse(dialog_id):
    token = request.headers.get("Authorization").split()
    if len(token) != 2:
        return get_data_error_result(message='Authorization is not valid!"')
    token = token[1]
    objs = APIToken.query(beta=token)
    if not objs:
        return get_data_error_result(message='Authentication error: API key is invalid!"')
    try:
        e, conv = DialogService.get_by_id(dialog_id)
        if not e:
            return get_data_error_result(message="Dialog not found!")
        conv = conv.to_dict()
        conv["avatar"] = conv["icon"]
        del conv["icon"]
        return get_json_result(data=conv)
    except Exception as e:
        return server_error_response(e)


@manager.route("/rm", methods=["POST"])  # type: ignore # type: ignore # noqa: F821
@login_required
def rm():
    conv_ids = request.json["conversation_ids"]
    try:
        for cid in conv_ids:
            exist, conv = ConversationService.get_by_id(cid)
            if not exist:
                return get_data_error_result(message="Conversation not found!")
            tenants = UserTenantService.query(user_id=current_user.id)
            for tenant in tenants:
                if DialogService.query(tenant_id=tenant.tenant_id, id=conv.dialog_id):
                    break
            else:
                return get_json_result(data=False, message="Only owner of conversation authorized for this operation.", code=settings.RetCode.OPERATING_ERROR)
            ConversationService.delete_by_id(cid)
        return get_json_result(data=True)
    except Exception as e:
        return server_error_response(e)


@manager.route("/list", methods=["GET"])  # type: ignore # noqa: F821
@login_required
def list_convsersation():
    dialog_id = request.args["dialog_id"]
    try:
        if not DialogService.query(tenant_id=current_user.id, id=dialog_id):
            return get_json_result(data=False, message="Only owner of dialog authorized for this operation.", code=settings.RetCode.OPERATING_ERROR)
        convs = ConversationService.query(dialog_id=dialog_id, order_by=ConversationService.model.create_time, reverse=True)

        convs = [d.to_dict() for d in convs]
        return get_json_result(data=convs)
    except Exception as e:
        return server_error_response(e)


@manager.route("/completion", methods=["POST"])  # type: ignore # noqa: F821
@login_required
@validate_request("conversation_id", "messages")
def completion():
    req = request.json
    msg = []
    temp_file_contents = []  # 存储临时文件内容
    
    for m in req["messages"]:
        if m["role"] == "system":
            continue
        if m["role"] == "assistant" and not msg:
            continue
        
        # 处理临时文件
        if m.get("temp_file_ids"):
            for file_id in m["temp_file_ids"]:
                try:
                    redis_key = f"temp_file:{file_id}"
                    file_data = REDIS_CONN.get(redis_key)
                    if file_data:
                        file_info = json.loads(file_data)
                        # 验证用户权限
                        if file_info.get('user_id') == current_user.id:
                            # 使用新的文件内容提取函数
                            file_content_bytes = base64.b64decode(file_info['content'])
                            content = extract_file_content(
                                file_content_bytes, 
                                file_info['filename'], 
                                file_info['content_type']
                            )
                            temp_file_contents.append({
                                'filename': file_info['filename'],
                                'content': content,
                                'content_type': file_info['content_type']
                            })
                except Exception as e:
                    print(f"Error processing temp file {file_id}: {e}")
        
        msg.append(m)
    
    # 如果有临时文件内容，将其添加到最后一条用户消息中
    if temp_file_contents and msg:
        last_user_msg = None
        for i in range(len(msg) - 1, -1, -1):
            if msg[i]["role"] == "user":
                last_user_msg = msg[i]
                break
        
        if last_user_msg:
            file_context = "\n\n[附件内容]:\n"
            for file_info in temp_file_contents:
                file_context += f"\n文件名: {file_info['filename']}\n内容:\n{file_info['content']}\n"
            last_user_msg["content"] += file_context
    
    message_id = msg[-1].get("id")
    try:
        e, conv = ConversationService.get_by_id(req["conversation_id"])
        if not e:
            return get_data_error_result(message="Conversation not found!")
        conv.message = deepcopy(req["messages"])
        e, dia = DialogService.get_by_id(conv.dialog_id)
        if not e:
            return get_data_error_result(message="Dialog not found!")
        del req["conversation_id"]
        del req["messages"]

        if not conv.reference:
            conv.reference = []
        else:

            def get_value(d, k1, k2):
                return d.get(k1, d.get(k2))

            for ref in conv.reference:
                if isinstance(ref, list):
                    continue
                ref["chunks"] = [
                    {
                        "id": get_value(ck, "chunk_id", "id"),
                        "content": get_value(ck, "content", "content_with_weight"),
                        "document_id": get_value(ck, "doc_id", "document_id"),
                        "document_name": get_value(ck, "docnm_kwd", "document_name"),
                        "dataset_id": get_value(ck, "kb_id", "dataset_id"),
                        "image_id": get_value(ck, "image_id", "img_id"),
                        "positions": get_value(ck, "positions", "position_int"),
                    }
                    for ck in ref.get("chunks", [])
                ]

        if not conv.reference:
            conv.reference = []
        conv.reference.append({"chunks": [], "doc_aggs": []})

        def stream():
            nonlocal dia, msg, req, conv
            try:
                for ans in chat(dia, msg, True, **req):
                    ans = structure_answer(conv, ans, message_id, conv.id)
                    yield "data:" + json.dumps({"code": 0, "message": "", "data": ans}, ensure_ascii=False) + "\n\n"
                ConversationService.update_by_id(conv.id, conv.to_dict())
            except Exception as e:
                traceback.print_exc()
                yield "data:" + json.dumps({"code": 500, "message": str(e), "data": {"answer": "**ERROR**: " + str(e), "reference": []}}, ensure_ascii=False) + "\n\n"
            yield "data:" + json.dumps({"code": 0, "message": "", "data": True}, ensure_ascii=False) + "\n\n"

        if req.get("stream", True):
            resp = Response(stream(), mimetype="text/event-stream")
            resp.headers.add_header("Cache-control", "no-cache")
            resp.headers.add_header("Connection", "keep-alive")
            resp.headers.add_header("X-Accel-Buffering", "no")
            resp.headers.add_header("Content-Type", "text/event-stream; charset=utf-8")
            return resp

        else:
            answer = None
            for ans in chat(dia, msg, **req):
                answer = structure_answer(conv, ans, message_id, req["conversation_id"])
                ConversationService.update_by_id(conv.id, conv.to_dict())
                break
            return get_json_result(data=answer)
    except Exception as e:
        return server_error_response(e)


# 用于文档撰写模式的问答调用
@manager.route("/writechat", methods=["POST"])  # type: ignore # noqa: F821
@login_required
@validate_request("question")
def writechat():
    req = request.json
    uid = current_user.id

    def stream():
        nonlocal req, uid
        try:
            # 获取 kb_ids，如果不存在则使用空数组
            kb_ids = req.get("kb_ids", [])
            for ans in write_dialog(req["question"], kb_ids, uid, req.get("similarity_threshold", 0.2), req.get("keyword_similarity_weight", 0.7), req.get("temperature", 1.0)):
                yield "data:" + json.dumps({"code": 0, "message": "", "data": ans}, ensure_ascii=False) + "\n\n"
        except Exception as e:
            yield "data:" + json.dumps({"code": 500, "message": str(e), "data": {"answer": "**ERROR**: " + str(e), "reference": []}}, ensure_ascii=False) + "\n\n"
        yield "data:" + json.dumps({"code": 0, "message": "", "data": True}, ensure_ascii=False) + "\n\n"

    resp = Response(stream(), mimetype="text/event-stream")
    resp.headers.add_header("Cache-control", "no-cache")
    resp.headers.add_header("Connection", "keep-alive")
    resp.headers.add_header("X-Accel-Buffering", "no")
    resp.headers.add_header("Content-Type", "text/event-stream; charset=utf-8")
    return resp

@manager.route("/uploadimage", methods=["POST"])  # type: ignore # noqa: F821
def uploadimage():
    if 'file' not in request.files:
        return jsonify({'error': '未检测到文件'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400
    url, err = upload_image(file)
    if err:
        return jsonify({'error': err}), 400
    return jsonify({'url': url})


@manager.route("/upload_temp_file", methods=["POST"])  # type: ignore # noqa: F821
def upload_temp_file():
    """上传临时文件到Redis，用于聊天问答"""
    try:
        # 手动检查认证
        jwt = Serializer(secret_key=settings.SECRET_KEY)
        authorization = request.headers.get("Authorization")
        
        if not authorization:
            return get_json_result(
                data=False, 
                message='No authorization.',
                code=settings.RetCode.AUTHENTICATION_ERROR
            )
        
        try:
            access_token = str(jwt.loads(authorization))
            
            user = UserService.query(
                access_token=access_token, status=StatusEnum.VALID.value
            )
            
            if not user:
                return get_json_result(
                    data=False, 
                    message='Invalid authorization.',
                    code=settings.RetCode.AUTHENTICATION_ERROR
                )
                
            current_user = user[0]
            
        except Exception as e:
            return get_json_result(
                data=False, 
                message='Invalid authorization.',
                code=settings.RetCode.AUTHENTICATION_ERROR
            )
        if 'file' not in request.files:
            return get_data_error_result(message="未检测到文件")
        
        file = request.files['file']
        
        if file.filename == '':
            return get_data_error_result(message="未选择文件")
        
        # 生成唯一的文件ID
        file_id = str(uuid.uuid4())
        
        # 读取文件内容
        file_content = file.read()
        
        # 检查文件大小（限制为10MB）
        if len(file_content) > 10 * 1024 * 1024:
            return get_data_error_result(message="文件大小不能超过10MB")
        
        # 对于文本文件，尝试解码内容以验证
        try:
            if file.content_type and 'text' in file.content_type:
                text_content = file_content.decode('utf-8')
        except Exception as e:
            pass
        
        # 将文件信息存储到Redis
        file_info = {
            'id': file_id,
            'filename': file.filename,
            'content_type': file.content_type or 'application/octet-stream',
            'size': len(file_content),
            'content': base64.b64encode(file_content).decode('utf-8'),
            'user_id': current_user.id,
            'conversation_id': request.form.get('conversation_id', ''),
            'upload_time': time.time()
        }
        
        # 存储到Redis，设置过期时间为1小时
        redis_key = f"temp_file:{file_id}"
        success = REDIS_CONN.set(redis_key, json.dumps(file_info, ensure_ascii=False), 3600)
        
        if not success:
            return get_data_error_result(message="存储文件失败，请稍后重试")
        
        # 验证存储是否成功
        stored_data = REDIS_CONN.get(redis_key)
        
        result_data = {
            'file_id': file_id,
            'filename': file.filename,
            'size': len(file_content),
            'content_type': file.content_type or 'application/octet-stream'
        }
        
        return get_json_result(data=result_data)
        
    except Exception as e:
        import traceback
        print(f"[ERROR] Upload temp file exception: {e}")
        traceback.print_exc()
        return server_error_response(e)


@manager.route("/get_temp_file/<file_id>", methods=["GET"])  # type: ignore # noqa: F821
def get_temp_file(file_id):
    """获取临时文件信息"""
    try:
        # 手动检查认证
        jwt = Serializer(secret_key=settings.SECRET_KEY)
        authorization = request.headers.get("Authorization")
        print(f"[DEBUG] get_temp_file - Authorization header: {authorization}")
        
        if not authorization:
            print("[DEBUG] get_temp_file - No Authorization header")
            return get_json_result(
                data=False, 
                message='No authorization.',
                code=settings.RetCode.AUTHENTICATION_ERROR
            )
        
        try:
            access_token = str(jwt.loads(authorization))
            print(f"[DEBUG] get_temp_file - Deserialized access_token: {access_token}")
            
            user = UserService.query(
                access_token=access_token, status=StatusEnum.VALID.value
            )
            
            if not user:
                print(f"[DEBUG] get_temp_file - No user found for access_token: {access_token}")
                return get_json_result(
                    data=False, 
                    message='Invalid authorization.',
                    code=settings.RetCode.AUTHENTICATION_ERROR
                )
                
            current_user = user[0]
            print(f"[DEBUG] get_temp_file - User authenticated: {current_user.email}")
            
        except Exception as e:
            print(f"[DEBUG] get_temp_file - Auth exception: {e}")
            return get_json_result(
                data=False, 
                message='Invalid authorization.',
                code=settings.RetCode.AUTHENTICATION_ERROR
            )
        redis_key = f"temp_file:{file_id}"
        file_data = REDIS_CONN.get(redis_key)
        
        if not file_data:
            return get_data_error_result(message="文件不存在或已过期")
        
        file_info = json.loads(file_data)
        
        # 验证用户权限
        if file_info.get('user_id') != current_user.id:
            return get_data_error_result(message="无权访问此文件")
        
        # 返回不包含content的文件信息
        return get_json_result(data={
            'id': file_info.get('id'),
            'filename': file_info.get('filename'),
            'content_type': file_info.get('content_type'),
            'size': file_info.get('size'),
            'upload_time': file_info.get('upload_time')
        })
        
    except Exception as e:
        return server_error_response(e)


@manager.route("/tts", methods=["POST"])  # type: ignore # noqa: F821
@login_required
def tts():
    req = request.json
    text = req["text"]

    tenants = TenantService.get_info_by(current_user.id)
    if not tenants:
        return get_data_error_result(message="Tenant not found!")

    tts_id = tenants[0]["tts_id"]
    if not tts_id:
        return get_data_error_result(message="No default TTS model is set")

    tts_mdl = LLMBundle(tenants[0]["tenant_id"], LLMType.TTS, tts_id)

    def stream_audio():
        try:
            for txt in re.split(r"[，。/《》？；：！\n\r:;]+", text):
                for chunk in tts_mdl.tts(txt):
                    yield chunk
        except Exception as e:
            yield ("data:" + json.dumps({"code": 500, "message": str(e), "data": {"answer": "**ERROR**: " + str(e)}}, ensure_ascii=False)).encode("utf-8")

    resp = Response(stream_audio(), mimetype="audio/mpeg")
    resp.headers.add_header("Cache-Control", "no-cache")
    resp.headers.add_header("Connection", "keep-alive")
    resp.headers.add_header("X-Accel-Buffering", "no")

    return resp


@manager.route("/delete_msg", methods=["POST"])  # type: ignore # noqa: F821
@login_required
@validate_request("conversation_id", "message_id")
def delete_msg():
    req = request.json
    e, conv = ConversationService.get_by_id(req["conversation_id"])
    if not e:
        return get_data_error_result(message="Conversation not found!")

    conv = conv.to_dict()
    for i, msg in enumerate(conv["message"]):
        if req["message_id"] != msg.get("id", ""):
            continue
        assert conv["message"][i + 1]["id"] == req["message_id"]
        conv["message"].pop(i)
        conv["message"].pop(i)
        conv["reference"].pop(max(0, i // 2 - 1))
        break

    ConversationService.update_by_id(conv["id"], conv)
    return get_json_result(data=conv)


@manager.route("/thumbup", methods=["POST"])  # type: ignore # noqa: F821
@login_required
@validate_request("conversation_id", "message_id")
def thumbup():
    req = request.json
    e, conv = ConversationService.get_by_id(req["conversation_id"])
    if not e:
        return get_data_error_result(message="Conversation not found!")
    up_down = req.get("set")
    feedback = req.get("feedback", "")
    conv = conv.to_dict()
    for i, msg in enumerate(conv["message"]):
        if req["message_id"] == msg.get("id", "") and msg.get("role", "") == "assistant":
            if up_down:
                msg["thumbup"] = True
                if "feedback" in msg:
                    del msg["feedback"]
            else:
                msg["thumbup"] = False
                if feedback:
                    msg["feedback"] = feedback
            break

    ConversationService.update_by_id(conv["id"], conv)
    return get_json_result(data=conv)


@manager.route("/ask", methods=["POST"])  # type: ignore # noqa: F821
@login_required
@validate_request("question", "kb_ids")
def ask_about():
    req = request.json
    uid = current_user.id

    def stream():
        nonlocal req, uid
        try:
            for ans in ask(req["question"], req["kb_ids"], uid):
                yield "data:" + json.dumps({"code": 0, "message": "", "data": ans}, ensure_ascii=False) + "\n\n"
        except Exception as e:
            yield "data:" + json.dumps({"code": 500, "message": str(e), "data": {"answer": "**ERROR**: " + str(e), "reference": []}}, ensure_ascii=False) + "\n\n"
        yield "data:" + json.dumps({"code": 0, "message": "", "data": True}, ensure_ascii=False) + "\n\n"

    resp = Response(stream(), mimetype="text/event-stream")
    resp.headers.add_header("Cache-control", "no-cache")
    resp.headers.add_header("Connection", "keep-alive")
    resp.headers.add_header("X-Accel-Buffering", "no")
    resp.headers.add_header("Content-Type", "text/event-stream; charset=utf-8")
    return resp


@manager.route("/mindmap", methods=["POST"])  # type: ignore # noqa: F821
@login_required
@validate_request("question", "kb_ids")
def mindmap():
    req = request.json
    kb_ids = req["kb_ids"]
    e, kb = KnowledgebaseService.get_by_id(kb_ids[0])
    if not e:
        return get_data_error_result(message="Knowledgebase not found!")

    embd_mdl = LLMBundle(kb.tenant_id, LLMType.EMBEDDING, llm_name=kb.embd_id)
    chat_mdl = LLMBundle(current_user.id, LLMType.CHAT)
    question = req["question"]
    ranks = settings.retrievaler.retrieval(question, embd_mdl, kb.tenant_id, kb_ids, 1, 12, 0.3, 0.3, aggs=False, rank_feature=label_question(question, [kb]))
    mindmap = MindMapExtractor(chat_mdl)
    mind_map = trio.run(mindmap, [c["content_with_weight"] for c in ranks["chunks"]])
    mind_map = mind_map.output
    if "error" in mind_map:
        return server_error_response(Exception(mind_map["error"]))
    return get_json_result(data=mind_map)


@manager.route("/related_questions", methods=["POST"])  # type: ignore # noqa: F821
@login_required
@validate_request("question")
def related_questions():
    req = request.json
    question = req["question"]
    chat_mdl = LLMBundle(current_user.id, LLMType.CHAT)
    prompt = """
Objective: To generate search terms related to the user's search keywords, helping users find more valuable information.
Instructions:
 - Based on the keywords provided by the user, generate 5-10 related search terms.
 - Each search term should be directly or indirectly related to the keyword, guiding the user to find more valuable information.
 - Use common, general terms as much as possible, avoiding obscure words or technical jargon.
 - Keep the term length between 2-4 words, concise and clear.
 - DO NOT translate, use the language of the original keywords.

### Example:
Keywords: Chinese football
Related search terms:
1. Current status of Chinese football
2. Reform of Chinese football
3. Youth training of Chinese football
4. Chinese football in the Asian Cup
5. Chinese football in the World Cup

Reason:
 - When searching, users often only use one or two keywords, making it difficult to fully express their information needs.
 - Generating related search terms can help users dig deeper into relevant information and improve search efficiency. 
 - At the same time, related terms can also help search engines better understand user needs and return more accurate search results.
 
"""
    ans = chat_mdl.chat(
        prompt,
        [
            {
                "role": "user",
                "content": f"""
Keywords: {question}
Related search terms:
    """,
            }
        ],
        {"temperature": 0.9},
    )
    return get_json_result(data=[re.sub(r"^[0-9]\. ", "", a) for a in ans.split("\n") if re.match(r"^[0-9]\. ", a)])
