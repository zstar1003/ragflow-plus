#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

from docx import Document
import re
import pandas as pd
from collections import Counter
from rag.nlp import rag_tokenizer
from io import BytesIO


class RAGFlowDocxParser:
    """
    Word文档(.docx)解析器，用于提取文档中的文本内容和表格。
    
    该解析器能够：
    1. 按页面范围提取文档中的段落文本及其样式
    2. 识别文档中的表格并将其转换为结构化文本
    3. 智能处理表格头部和内容，生成语义化的文本描述
    """

    def __extract_table_content(self, tb):
        """
        从Word表格对象中提取内容并转换为DataFrame
        
        参数:
            tb: docx库的Table对象
            
        返回:
            处理后的表格内容文本列表
        """
        df = []
        for row in tb.rows:
            df.append([c.text for c in row.cells])
        return self.__compose_table_content(pd.DataFrame(df))

    def __compose_table_content(self, df):
        """
        将表格DataFrame转换为语义化的文本描述
        
        通过识别表格的结构特征(如表头、数据类型等)，将表格转换为更易于理解的文本形式
        
        参数:
            df: 包含表格内容的DataFrame
            
        返回:
            表格内容的文本表示列表
        """

        def blockType(b):
            """
            识别单元格内容的类型
            
            通过正则表达式和文本特征分析，将单元格内容分类为不同类型：
            - Dt: 日期类型
            - Nu: 数字类型
            - Ca: 代码/ID类型
            - En: 英文文本
            - NE: 数字和文本混合
            - Sg: 单字符
            - Tx: 短文本
            - Lx: 长文本
            - Nr: 人名
            - Ot: 其他类型
            
            参数:
                b: 单元格文本内容
                
            返回:
                内容类型的字符串标识
            """
            patt = [
                ("^(20|19)[0-9]{2}[年/-][0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^(20|19)[0-9]{2}年$", "Dt"),
                (r"^(20|19)[0-9]{2}[年/-][0-9]{1,2}月*$", "Dt"),
                ("^[0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^第*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}年*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}[ABCDE]$", "DT"),
                ("^[0-9.,+%/ -]+$", "Nu"),
                (r"^[0-9A-Z/\._~-]+$", "Ca"),
                (r"^[A-Z]*[a-z' -]+$", "En"),
                (r"^[0-9.,+-]+[0-9A-Za-z/$￥%<>（）()' -]+$", "NE"),
                (r"^.{1}$", "Sg")
            ]
            for p, n in patt:
                if re.search(p, b):
                    return n
            tks = [t for t in rag_tokenizer.tokenize(b).split() if len(t) > 1]
            if len(tks) > 3:
                if len(tks) < 12:
                    return "Tx"
                else:
                    return "Lx"

            if len(tks) == 1 and rag_tokenizer.tag(tks[0]) == "nr":
                return "Nr"

            return "Ot"

        # 表格至少需要两行才能处理
        if len(df) < 2:
            return []
            
        # 统计表格中最常见的内容类型
        max_type = Counter([blockType(str(df.iloc[i, j])) for i in range(
            1, len(df)) for j in range(len(df.iloc[i, :]))])
        max_type = max(max_type.items(), key=lambda x: x[1])[0]

        # 获取表格列数
        colnm = len(df.iloc[0, :])
        # 默认第一行为表头
        hdrows = [0]  # 表头不一定出现在第一行
        
        # 如果表格主要是数字类型，则识别非数字行作为表头
        if max_type == "Nu":
            for r in range(1, len(df)):
                tys = Counter([blockType(str(df.iloc[r, j]))
                              for j in range(len(df.iloc[r, :]))])
                tys = max(tys.items(), key=lambda x: x[1])[0]
                if tys != max_type:
                    hdrows.append(r)

        # 处理表格内容，将每行转换为文本
        lines = []
        for i in range(1, len(df)):
            # 跳过表头行
            if i in hdrows:
                continue
                
            # 计算当前行之前的表头行
            hr = [r - i for r in hdrows]
            hr = [r for r in hr if r < 0]
            
            # 找到最近的连续表头行
            t = len(hr) - 1
            while t > 0:
                if hr[t] - hr[t - 1] > 1:
                    hr = hr[t:]
                    break
                t -= 1
                
            # 为每列构建表头描述
            headers = []
            for j in range(len(df.iloc[i, :])):
                t = []
                for h in hr:
                    x = str(df.iloc[i + h, j]).strip()
                    if x in t:
                        continue
                    t.append(x)
                t = ",".join(t)
                if t:
                    t += ": "
                headers.append(t)
                
            # 构建每行的文本表示
            cells = []
            for j in range(len(df.iloc[i, :])):
                if not str(df.iloc[i, j]):
                    continue
                cells.append(headers[j] + str(df.iloc[i, j]))
            lines.append(";".join(cells))

        # 根据列数决定返回格式
        if colnm > 3:
            return lines
        return ["\n".join(lines)]

    def __call__(self, fnm, from_page=0, to_page=100000000):
        """
        解析Word文档，提取指定页面范围内的文本和表格
        
        参数:
            fnm: 文件名或二进制内容
            from_page: 起始页码(从0开始)
            to_page: 结束页码
            
        返回:
            元组(secs, tbls)，其中:
            - secs: 段落内容列表，每项为(文本内容, 样式名称)的元组
            - tbls: 表格内容列表
        """
        # 根据输入类型创建Document对象
        self.doc = Document(fnm) if isinstance(
            fnm, str) else Document(BytesIO(fnm))
        pn = 0 # 当前解析页码
        secs = [] # 存储解析的段落内容
        
        # 遍历文档中的所有段落
        for p in self.doc.paragraphs:
            # 如果超出指定页码范围，停止解析
            if pn > to_page:
                break

            runs_within_single_paragraph = [] # 保存在页面范围内的文本片段
            # 遍历段落中的所有文本片段(run)
            for run in p.runs:
                if pn > to_page:
                    break
                # 如果当前页码在指定范围内且段落有内容，则添加文本
                if from_page <= pn < to_page and p.text.strip():
                    runs_within_single_paragraph.append(run.text) # 先添加文本片段

                # 检查页面分隔符
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1

            # 将段落文本和样式添加到结果列表
            secs.append(("".join(runs_within_single_paragraph), p.style.name if hasattr(p.style, 'name') else '')) # 然后将文本片段连接为段落的一部分

        # 提取所有表格内容
        tbls = [self.__extract_table_content(tb) for tb in self.doc.tables]
        return secs, tbls
